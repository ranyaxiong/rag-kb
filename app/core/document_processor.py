"""
文档处理核心模块
"""
import os
import uuid
from datetime import datetime
from typing import List, Dict, Any
import logging
import shutil

from langchain_community.document_loaders import (
    PyPDFLoader,
    UnstructuredWordDocumentLoader, 
    TextLoader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from app.core.config import settings

logger = logging.getLogger(__name__)


class MarkdownLoader:
    """简单的Markdown文档加载器"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """加载markdown文件内容"""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 创建文档对象
            doc = Document(
                page_content=content,
                metadata={'source': self.file_path}
            )
            
            return [doc]
        except Exception as e:
            logger.error(f"Error loading markdown file {self.file_path}: {str(e)}")
            raise


class WordDocumentLoader:
    """自定义Word文档加载器"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """加载Word文档内容"""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            content = ""
            
            # 检查文件扩展名
            if self.file_path.lower().endswith('.docx'):
                content = self._extract_docx_content()
            elif self.file_path.lower().endswith('.doc'):
                content = self._extract_doc_content()
            else:
                raise ValueError(f"不支持的文件格式: {self.file_path}")
            
            # 创建文档对象
            doc = Document(
                page_content=content,
                metadata={'source': self.file_path}
            )
            
            return [doc]
        except Exception as e:
            logger.error(f"Error loading Word document {self.file_path}: {str(e)}")
            raise
    
    def _extract_docx_content(self) -> str:
        """从docx文件提取文本内容"""
        import zipfile
        import xml.etree.ElementTree as ET
        
        try:
            with zipfile.ZipFile(self.file_path, 'r') as docx:
                # 读取document.xml文件
                document_xml = docx.read('word/document.xml')
                
                # 解析XML
                root = ET.fromstring(document_xml)
                
                # Word文档的命名空间
                namespace = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                # 提取所有文本节点
                text_elements = root.findall('.//w:t', namespace)
                content = ''.join([elem.text or '' for elem in text_elements])
                
                return content
                
        except Exception as e:
            logger.error(f"Error extracting content from DOCX file: {str(e)}")
            # 如果XML解析失败，尝试使用python-docx作为备选
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(self.file_path)
                return "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except ImportError:
                raise ValueError(f"无法解析Word文档，请确保文件格式正确: {str(e)}")
            except Exception as docx_error:
                raise ValueError(f"无法读取Word文档内容: {str(docx_error)}")
    
    def _extract_doc_content(self) -> str:
        """从.doc文件提取文本内容"""
        try:
            # 尝试基础文本提取方法（从二进制中提取可读文本）
            try:
                logger.info("尝试从.doc文件中提取文本内容")
                
                with open(self.file_path, 'rb') as f:
                    raw_content = f.read()
                
                # 方法1：提取连续的可打印字符
                import re
                
                # 查找连续的可打印字符（包括中文Unicode范围）
                # ASCII可打印字符
                ascii_parts = re.findall(rb'[\x20-\x7E]{4,}', raw_content)
                ascii_text = ' '.join([part.decode('ascii', errors='ignore') for part in ascii_parts])
                
                # 尝试UTF-16编码（Word常用）
                utf16_text = ""
                try:
                    # 去除BOM标记并尝试解码
                    if raw_content.startswith(b'\xff\xfe') or raw_content.startswith(b'\xfe\xff'):
                        utf16_text = raw_content[2:].decode('utf-16', errors='ignore')
                    else:
                        utf16_text = raw_content.decode('utf-16le', errors='ignore')
                    
                    # 清理UTF-16文本，去除控制字符
                    utf16_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', utf16_text)
                    utf16_text = ' '.join(utf16_text.split())  # 规范化空白字符
                    
                except Exception as e:
                    logger.debug(f"UTF-16解码失败: {str(e)}")
                
                # 选择最好的结果
                content = ""
                if len(utf16_text.strip()) > len(ascii_text.strip()):
                    content = utf16_text.strip()
                else:
                    content = ascii_text.strip()
                
                # 检查是否提取到有意义的内容
                if content and len(content) >= 10:
                    logger.warning(f"从.doc文件中提取到{len(content)}个字符的内容（可能不完整）")
                    return content
                else:
                    logger.warning("未能从.doc文件中提取到足够的文本内容")
                    
            except Exception as e:
                logger.error(f"文本提取过程中出错: {str(e)}")
            
            # 如果无法提取内容，提供清晰的错误信息和建议
            raise ValueError(
                "❌ 无法处理.doc格式文件\n\n"
                "💡 解决方案:\n"
                "1. 【推荐】使用Microsoft Word打开文件，另存为.docx格式\n"
                "2. 【简单】将文件另存为.txt格式\n"
                "3. 【在线转换】使用在线转换工具将.doc转为.docx\n\n"
                "📋 技术说明:\n"
                ".doc是Microsoft Word 97-2003的二进制格式，需要专门的解析库。\n"
                "系统目前完全支持.docx、.pdf、.txt、.md等格式。"
            )
            
        except ValueError:
            # 重新抛出我们的用户友好错误
            raise
        except Exception as e:
            logger.error(f"Error extracting content from DOC file: {str(e)}")
            raise ValueError(f"处理.doc文件时发生错误: {str(e)}")


class SmartPDFLoader:
    """智能PDF加载器，自动选择最佳处理策略"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        # 惰性导入，避免模块导入失败导致整个处理器不可用
        try:
            from app.core.enhanced_pdf_processor import EnhancedPDFProcessor
            self.enhanced_processor = EnhancedPDFProcessor()
        except Exception as e:
            logger.warning(f"EnhancedPDFProcessor unavailable, fallback to PyPDFLoader: {e}")
            self.enhanced_processor = None
        
    def load(self) -> List[Document]:
        """加载PDF文档，智能选择处理方法"""
        try:
            # 检查是否可以使用增强处理器
            processing_info = None
            if self.enhanced_processor is not None:
                processing_info = self.enhanced_processor.get_processing_info(self.file_path)
            
            # 使用增强处理器
            if self.enhanced_processor is not None:
                try:
                    documents = self.enhanced_processor.load(self.file_path)
                    if documents and any(doc.page_content.strip() for doc in documents):
                        logger.info(f"使用增强PDF处理器成功处理: {self.file_path}")
                        return documents
                except Exception as e:
                    logger.warning(f"增强PDF处理器失败，回退到标准处理器: {str(e)}")
            
            # 回退到标准PDF处理器
            standard_loader = PyPDFLoader(self.file_path)
            documents = standard_loader.load()
            
            if not documents or not any(doc.page_content.strip() for doc in documents):
                logger.warning(f"标准PDF处理器也未提取到内容: {self.file_path}")
                
                # 最后尝试：提供处理建议
                suggestions = self._get_processing_suggestions(processing_info)
                error_msg = f"无法从PDF中提取文本内容。\n\n{suggestions}"
                raise ValueError(error_msg)
            
            logger.info(f"使用标准PDF处理器处理: {self.file_path}")
            return documents
            
        except Exception as e:
            logger.error(f"PDF处理完全失败: {self.file_path} - {str(e)}")
            raise
    
    def _get_processing_suggestions(self, processing_info: Dict) -> str:
        """基于PDF分析结果提供处理建议"""
        suggestions = []
        
        pdf_analysis = processing_info.get('pdf_analysis', {})
        is_scanned = pdf_analysis.get('is_scanned', False)
        ocr_available = processing_info.get('ocr_available', False)
        
        if is_scanned:
            suggestions.append("📋 检测结果: 这似乎是一个扫描类PDF（图像格式）")
            
            if not ocr_available:
                suggestions.extend([
                    "💡 解决方案:",
                    "1. 安装OCR依赖以自动处理扫描类PDF:",
                    "   pip install pytesseract pillow",
                    "2. 安装Tesseract-OCR引擎:",
                    "   - Windows: https://github.com/UB-Mannheim/tesseract/wiki",
                    "   - 或将PDF转换为可搜索的PDF格式",
                ])
            else:
                suggestions.append("❗ OCR功能可用但处理失败，可能是PDF格式问题")
        else:
            suggestions.extend([
                "📋 这个PDF应该包含可提取的文本，但提取失败",
                "💡 建议:",
                "1. 检查PDF是否受密码保护",
                "2. 尝试用其他PDF查看器打开验证文件完整性",
                "3. 将PDF另存为新文件后重试"
            ])
        
        return "\n".join(suggestions)


class DocumentProcessor:
    """文档处理核心类"""
    
    def __init__(self):
        self.loaders = {
            '.pdf': SmartPDFLoader,
            '.docx': WordDocumentLoader,
            '.doc': WordDocumentLoader,
            '.txt': TextLoader,
            '.md': MarkdownLoader
        }
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
        
        self.supported_extensions = set(self.loaders.keys())
    
    def is_supported_file(self, filename: str) -> bool:
        """检查文件是否支持"""
        _, ext = os.path.splitext(filename.lower())
        return ext in self.supported_extensions
    
    def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """保存上传的文件"""
        try:
            # 生成唯一文件名
            file_id = str(uuid.uuid4())
            _, ext = os.path.splitext(filename)
            safe_filename = f"{file_id}_{filename}"
            
            # 创建日期目录
            date_dir = datetime.now().strftime("%Y-%m-%d")
            full_dir = os.path.join(settings.upload_dir, date_dir)
            os.makedirs(full_dir, exist_ok=True)
            
            file_path = os.path.join(full_dir, safe_filename)
            
            # 保存文件
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"File saved: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error saving file {filename}: {str(e)}")
            raise
    
    def save_to_temp_file(self, file_content: bytes, filename: str) -> str:
        """先将上传内容保存到容器本地临时目录，写入更快，返回临时路径"""
        try:
            file_id = str(uuid.uuid4())
            safe_filename = f"{file_id}_{filename}"
            date_dir = datetime.now().strftime("%Y-%m-%d")
            full_dir = os.path.join(settings.temp_upload_dir, date_dir)
            os.makedirs(full_dir, exist_ok=True)
            temp_path = os.path.join(full_dir, safe_filename)
            with open(temp_path, 'wb') as f:
                f.write(file_content)
            logger.info(f"Temp file saved: {temp_path}")
            return temp_path
        except Exception as e:
            logger.error(f"Error saving temp file {filename}: {str(e)}")
            raise

    def move_to_upload_dir(self, temp_file_path: str, original_filename: str) -> str:
        """将临时文件移动到最终上传目录，保持同名，返回最终路径"""
        try:
            if not os.path.exists(temp_file_path):
                raise FileNotFoundError(f"Temp file not found: {temp_file_path}")
            base_name = os.path.basename(temp_file_path)
            date_dir = datetime.now().strftime("%Y-%m-%d")
            dest_dir = os.path.join(settings.upload_dir, date_dir)
            os.makedirs(dest_dir, exist_ok=True)
            dest_path = os.path.join(dest_dir, base_name)
            # 如果目标已存在（极少见），在文件名后追加后缀
            if os.path.exists(dest_path):
                name, ext = os.path.splitext(base_name)
                dest_path = os.path.join(dest_dir, f"{name}_moved{ext}")
            shutil.move(temp_file_path, dest_path)
            logger.info(f"Moved temp file to upload dir: {dest_path}")
            return dest_path
        except Exception as e:
            logger.error(f"Error moving temp file to upload dir: {str(e)}")
            raise

    def is_in_temp_dir(self, file_path: str) -> bool:
        """判断路径是否位于临时上传目录下"""
        try:
            return os.path.abspath(file_path).startswith(os.path.abspath(settings.temp_upload_dir))
        except Exception:
            return False
    
    def load_document(self, file_path: str) -> List[Document]:
        """加载文档内容"""
        try:
            _, ext = os.path.splitext(file_path.lower())
            
            if ext not in self.loaders:
                raise ValueError(f"Unsupported file type: {ext}")
            
            loader_class = self.loaders[ext]
            loader = loader_class(file_path)
            
            documents = loader.load()
            logger.info(f"Loaded {len(documents)} document(s) from {file_path}")
            
            return documents
            
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {str(e)}")
            raise
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """分割文档为块"""
        try:
            chunks = self.text_splitter.split_documents(documents)
            logger.info(f"Split into {len(chunks)} chunks")
            
            # 添加块索引到元数据
            for i, chunk in enumerate(chunks):
                chunk.metadata['chunk_index'] = i
                chunk.metadata['chunk_id'] = str(uuid.uuid4())
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error splitting documents: {str(e)}")
            raise
    
    def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """处理单个文档的完整流程"""
        try:
            doc_id = str(uuid.uuid4())
            
            # 加载文档
            documents = self.load_document(file_path)
            
            # 添加文档元数据
            for doc in documents:
                doc.metadata.update({
                    'document_id': doc_id,
                    'filename': filename,
                    'file_path': file_path,
                    'processed_at': datetime.now().isoformat()
                })
            
            # 分割文档
            chunks = self.split_documents(documents)
            
            result = {
                'document_id': doc_id,
                'filename': filename,
                'file_path': file_path,
                'chunks': chunks,
                'chunk_count': len(chunks),
                'status': 'completed'
            }
            
            logger.info(f"Successfully processed document: {filename}")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            return {
                'document_id': str(uuid.uuid4()),
                'filename': filename,
                'file_path': file_path,
                'chunks': [],
                'chunk_count': 0,
                'status': 'failed',
                'error_message': str(e)
            }
    
    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """获取文档基本信息"""
        try:
            stat = os.stat(file_path)
            filename = os.path.basename(file_path)
            _, ext = os.path.splitext(filename)
            
            return {
                'filename': filename,
                'file_type': ext.lower(),
                'file_size': stat.st_size,
                'upload_time': datetime.fromtimestamp(stat.st_mtime),
                'is_supported': self.is_supported_file(filename)
            }
        except Exception as e:
            logger.error(f"Error getting document info for {file_path}: {str(e)}")
            return None